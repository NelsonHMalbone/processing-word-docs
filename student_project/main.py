import docx
import os

directory = 'files'
output_file = 'files/merge_docs.docx'

# Generate filenames
filenames = os.listdir(directory)

# Generate filepaths
filepaths = [os.path.join(directory, filename) for filename in filenames
             if filename.endswith('.docx')] # just incase other files are inside of the same file

# create a new doc object
merge_doc = docx.Document()

'Iterate over all DOCX files in the specified folder'

for filepath in filepaths:
    documents = docx.Document(filepath)

    'Append each paragraph from the current document to the merged document'
    para = documents.paragraphs

    for para in documents.paragraphs:


        'Create a new paragraph in the merged document with the same formatting'

        new_para = merge_doc.add_paragraph()
        new_para._element.addprevious(para._element)
        # saving to new file
merge_doc.save(output_file)