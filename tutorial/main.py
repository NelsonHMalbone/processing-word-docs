"combining word documents to one document with python"
# want to insert paragraph from file 2 to file 1


# need to install with pip install python-docx
import docx

# our input is the two files in files folder
docpath1 = 'files/panda1.docx'
docpath2 = 'files/panda2.docx'

# creating a special class from the docx library
"need for however many docs you need to go thur"
doc1 = docx.Document(docpath1)
doc2 = docx.Document(docpath2)

"""
# read and print out the first file
paragraphs = doc1.paragraphs

# extracting text
para1 = paragraphs[0]  # first paragraph
para_text = para1.text  # the text from para1
"""

# get the text from from 2
"text object"
# text = doc2.paragraphs[0].text # here we just extract the text no font nothing else
"para object"
para = doc2.paragraphs[0] # this method preserves the format of text and such

# doc2.paragraphs[0] paragraph object
# .text is getting the text

# method 1 to get paragraph from file 2
'if using the para object we dont need this line'
#new_paragraph = doc1.add_paragraph(text) # get a list of paragraphs

# method 2 this method will place the new paragraph into the right location
paragraphs = doc1.paragraphs # the paragraphs of doc 1
"only for the text object"
# paragraphs[1]._element.addnext(new_paragraph._element) # access the 2nd paragraph and places the new section below it
"only for the para object"
paragraphs[1]._element.addnext(para._element)


# access the modified file
doc1.save("files/updated_panda.docx")
